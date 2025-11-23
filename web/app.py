import os
import uuid
import json
import tempfile
import logging
import datetime
import time
from typing import Any, Dict, List, Optional, Tuple

from flask import Flask, render_template, request, jsonify, send_file, session, redirect, url_for, make_response
from werkzeug.utils import secure_filename
import filetype
from functools import wraps
import re
import csv

# Additional Format Libs
import yaml
from fpdf import FPDF
from docx import Document as DocxDocument
from openpyxl import Workbook
import xml.etree.ElementTree as ET
from xml.dom import minidom

# Backend
from image_ocr import extract_text_from_path
from llm_model import (
    summarize_text,
    translate_text,
    organize_entities,
    is_llm_available
)
from logging_config import setup_logging
from web.security import security, require_rate_limit

# ----------------------------
# Environment Loading
# ----------------------------
def _read_text_any_encoding(path: str) -> str:
    with open(path, "rb") as f:
        data = f.read()
    for enc in ("utf-8-sig", "utf-16", "utf-16-le", "utf-16-be", "cp1252"):
        try:
            return data.decode(enc)
        except Exception:
            continue
    return data.decode("utf-8", errors="ignore")

def _strip_quotes(v: str) -> str:
    v = v.strip()
    if len(v) >= 2 and ((v[0] == v[-1]) and v[0] in ("'", '"')):
        return v[1:-1]
    return v

def _parse_env_line(line: str) -> Optional[Tuple[str, str]]:
    line = line.strip()
    if not line or line.startswith("#"):
        return None
    if line.lower().startswith("export "):
        line = line[7:].strip()
    sep = "=" if "=" in line else (":" if ":" in line else None)
    if not sep:
        return None
    k, v = line.split(sep, 1)
    k = k.strip()
    v = _strip_quotes(v.strip())
    if not k:
        return None
    return k, v

def _load_env_file(path: str, override: bool) -> None:
    if not os.path.isfile(path):
        return
    try:
        content = _read_text_any_encoding(path)
        for raw_line in content.splitlines():
            parsed = _parse_env_line(raw_line)
            if not parsed:
                continue
            k, v = parsed
            if not override and os.environ.get(k):
                continue
            os.environ[k] = v
    except Exception:
        pass

def _load_env_pair() -> None:
    here = os.path.dirname(os.path.abspath(__file__))
    root = os.path.dirname(here)
    for d in (root, here):
        _load_env_file(os.path.join(d, ".env"), override=True)

_load_env_pair()

setup_logging()
logger = logging.getLogger(__name__)

# ----------------------------
# Flask App
# ----------------------------
app = Flask(__name__, template_folder="templates")
app.config["UPLOAD_FOLDER"] = tempfile.gettempdir()
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024
app.secret_key = os.urandom(24)
app.permanent_session_lifetime = datetime.timedelta(minutes=30) # Auto-logout
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
# Secure cookie requires HTTPS, but we allow HTTP for now if user reverts.
# app.config["SESSION_COOKIE_SECURE"] = True 

app.security = security  # Attach for run_web.py access

ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg", "bmp", "tiff", "webp"}

def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def _safe_bool(value: Any, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        v = value.strip().lower()
        if v in ("true", "1", "yes", "y", "on"):
            return True
        if v in ("false", "0", "no", "n", "off"):
            return False
    return default

# ----------------------------
# Middleware
# ----------------------------
@app.after_request
def add_security_headers(response):
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY'
    # Content-Security-Policy: Allow CDNs
    csp = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net https://cdnjs.cloudflare.com; "
        "style-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net https://cdnjs.cloudflare.com https://fonts.googleapis.com; "
        "font-src 'self' https://cdnjs.cloudflare.com https://fonts.gstatic.com; "
        "img-src 'self' data: blob:;"
    )
    response.headers['Content-Security-Policy'] = csp
    return response

@app.before_request
def check_auth():
    ip = request.remote_addr or "127.0.0.1"
    if security.is_blacklisted(ip):
         return jsonify({"error": "Access Denied (Blacklisted)"}), 403

    if request.path.startswith('/static') or request.path == '/login' or request.path == '/api/health':
        return
    if request.remote_addr in ('127.0.0.1', '::1'):
        return
    if not session.get('authenticated'):
        return redirect(url_for('login_page'))

# ----------------------------
# Routes
# ----------------------------
@app.route("/login", methods=["GET"])
def login_page():
    if session.get('authenticated'):
        return redirect(url_for('index'))
    return render_template("login.html")

@app.route("/login", methods=["POST"])
@require_rate_limit(limit=5, window=60) # 5 attempts per minute
def do_login():
    data = request.json or {}
    code = str(data.get("code", "")).strip()
    auth_type = security.verify_token(code)
    
    if auth_type:
        session['authenticated'] = True
        session.permanent = (auth_type == "ADMIN")
        security.log("AUTH_SUCCESS", f"User logged in as {auth_type}", request.remote_addr)
        return jsonify({"success": True})
    
    security.record_violation(request.remote_addr)
    security.log("AUTH_FAIL", "Invalid access code", request.remote_addr)
    return jsonify({"success": False, "error": "Invalid or expired access code"}), 401

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({
        "status": "ok",
        "llm_available": is_llm_available(),
        "ocr_key_loaded": bool((os.environ.get("OCR_SPACE_API_KEY") or "").strip())
    }), 200

@app.route("/upload", methods=["POST"])
def upload_files():
    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "No files provided"}), 400
    saved_files = []
    for f in files:
        if not f or not f.filename or not allowed_file(f.filename):
            continue
        
        # Magic Byte Validation (Security Patch)
        kind = filetype.guess(f.stream)
        f.stream.seek(0) # Reset stream after reading
        if kind is None or kind.mime.split('/')[0] != 'image':
             # Fallback for some obscure types, or reject. Secure: Reject.
             logger.warning(f"Rejected file {f.filename} - Invalid MIME type: {kind.mime if kind else 'Unknown'}")
             continue

        ext = f.filename.rsplit(".", 1)[1].lower()
        unique_name = secure_filename(f"{uuid.uuid4().hex}.{ext}")
        path = os.path.join(app.config["UPLOAD_FOLDER"], unique_name)
        f.save(path)
        saved_files.append(path)
    if not saved_files:
        return jsonify({"error": "No valid image files uploaded (Check file headers)"}), 400
    return jsonify({"files": saved_files}), 200

def sanitize_prompt(text: str) -> str:
    # Basic Anti-Injection Filter
    if not text: return ""
    # Block common override attempts
    patterns = [r"ignore previous", r"system instruction", r"system override"]
    clean = text
    for p in patterns:
        clean = re.sub(p, "[FILTERED]", clean, flags=re.IGNORECASE)
    return clean

@app.route("/process", methods=["POST"])
@require_rate_limit(limit=10, window=60)
def process_images():
    """
    Main processing endpoint: OCR -> [Summarize] -> [Entity Order] -> [Translate]
    """
    data = request.json or {}
    files = data.get("files", [])
    ocr_mode = data.get("ocr_mode", "auto")
    translation = data.get("translation", "original") # original, ar, en
    summarize = data.get("summarize", False)
    entity_order = data.get("entity_order", []) # list of strings

    if not isinstance(files, list) or not files:
        return jsonify({"error": "No files to process"}), 400
    
    final_text_parts = []
    all_tables = []
    
    # Timing stats
    t_start_ocr = time.time()
    
    # ... OCR Loop ...
    for path in files:
        try:
            result = extract_text_from_path(path, mode=ocr_mode)
            if isinstance(result, dict):
                if result.get("error"):
                     continue
                text = (result or {}).get("text", "")
                if text:
                    final_text_parts.append(text)
                # Collect tables
                if result.get("tables"):
                    all_tables.extend(result["tables"])
        except Exception as e:
            logger.exception("OCR failed")
            
    ocr_duration = time.time() - t_start_ocr
    final_text = "\n\n".join(final_text_parts)

    logger.info(f"Processing text length: {len(final_text)}")
    
    # AI Processing
    t_start_llm = time.time()
    
    # 1. Summarization (if requested)
    if summarize:
        logger.info(">> Triggering Summarization...")
        final_text = summarize_text(final_text)

    # 2. Entity Organization (if requested and LLM available)
    if entity_order and isinstance(entity_order, list) and is_llm_available():
        # Sanitize Inputs
        safe_order = [sanitize_prompt(str(i)) for i in entity_order]
        logger.info(f">> Triggering Entity Reorder: {safe_order}")
        final_text = organize_entities(final_text, safe_order)

    # 3. Translation (if requested)
    if translation != "original" and translation in ("ar", "en"):
        logger.info(f">> Triggering Translation to {translation}")
        final_text = translate_text(final_text, translation)
        
    ai_duration = time.time() - t_start_llm
    
    # Sanitize AI Output (Security)
    final_text = security.sanitize_ai_output(final_text)
        
    return jsonify({
        "text": final_text,
        "tables": all_tables,
        "ocr_mode": ocr_mode,
        "timing": {
            "ocr": round(ocr_duration, 2),
            "ai": round(ai_duration, 2)
        }
    }), 200

@app.route("/api/reformat", methods=["POST"])
def reformat_text_route():
    try:
        data = request.json or {}
        text = data.get("text", "")
        order = data.get("entity_order", [])
        
        if not text or not order:
            return jsonify({"error": "Missing text or order list"}), 400
            
        if not is_llm_available():
            return jsonify({"error": "AI is unavailable"}), 503

        logger.info(f"Reformatting text with order: {order}")
        new_text = organize_entities(text, order)
        
        return jsonify({"text": new_text})
        
    except Exception as e:
        logger.exception("Reformat failed")
        return jsonify({"error": str(e)}), 500

@app.route("/save", methods=["POST"])
def save_text():
    data = request.json or {}
    text = data.get("text", "")
    filename = secure_filename(data.get("filename") or "output")
    fmt = (data.get("format") or "txt").lower()
    
    out_path = os.path.join(tempfile.gettempdir(), f"{filename}.{fmt}")
    
    try:
        # JSON
        if fmt == "json":
            with open(out_path, "w", encoding="utf-8") as f:
                json.dump({"text": text}, f, indent=2, ensure_ascii=False)
                
        # YAML
        elif fmt == "yaml":
            with open(out_path, "w", encoding="utf-8") as f:
                yaml.dump({"text": text}, f, allow_unicode=True)
                
        # XML
        elif fmt == "xml":
            root = ET.Element("root")
            # Split lines for structure, or just one block
            lines = text.split('\n')
            for line in lines:
                if line.strip():
                    child = ET.SubElement(root, "line")
                    child.text = line
            
            raw_xml = ET.tostring(root, encoding='utf-8')
            parsed = minidom.parseString(raw_xml)
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(parsed.toprettyxml(indent="  "))
                
        # HTML
        elif fmt == "html":
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(f"""<!DOCTYPE html>
<html>
<head><meta charset='UTF-8'><title>{filename}</title></head>
<body style='font-family: Arial; white-space: pre-wrap;'>
{text}
</body>
</html>""")

        # CSV
        elif fmt == "csv":
            with open(out_path, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f)
                writer.writerow(["Line Content"]) # Header
                for line in text.split('\n'):
                     writer.writerow([line])
                    
        # EXCEL (XLSX)
        elif fmt == "xlsx" or fmt == "excel":
            wb = Workbook()
            ws = wb.active
            ws.title = "Text Data"
            ws.append(["Line Content"])
            for line in text.split('\n'):
                ws.append([line])
            wb.save(out_path)
            
        # WORD (DOCX)
        elif fmt == "docx" or fmt == "word":
            doc = DocxDocument()
            doc.add_heading(filename, 0)
            doc.add_paragraph(text)
            doc.save(out_path)
            
        # PDF
        elif fmt == "pdf":
            # Basic PDF output (Warning: Arabic might need special font, but we start basic)
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            # FPDF has trouble with Unicode/Arabic without font setup. 
            # We will use latin-1 safe encoding or replace checks.
            # Ideally we embed a font, but for now we try basic replacement.
            # A proper Arabic PDF solution requires a font file like auto-downloading 'Amiri'.
            # We'll encode logically to avoid crash.           
            try:
                # Naive dump, usually works for English
                for line in text.split('\n'):
                    pdf.cell(200, 10, txt=line.encode('latin-1', 'replace').decode('latin-1'), ln=1, align='L')
            except:
                pass 
            pdf.output(out_path)

        # TXT (Default)
        else:
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(text)

        return send_file(out_path, as_attachment=True)
        
    except Exception as e:
        logger.error(f"Save failed: {e}")
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(host="127.0.0.1", port=5000, debug=False)
